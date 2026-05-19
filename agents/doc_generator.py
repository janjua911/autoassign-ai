"""
agents/doc_generator.py
Doc Generator Agent — creates formatted DOCX + PDF from verified answers.
"""

from pathlib import Path
from datetime import datetime
from loguru import logger
from core.database import db
from core.config import cfg

OUTPUT_DIR = Path(__file__).parent.parent / "outputs"


class DocGeneratorAgent:
    """
    Generates a professional assignment document (DOCX + PDF).
    Uses python-docx for Word, reportlab for PDF.
    """

    def run(self, assignment_id: int) -> tuple[str, str]:
        """
        Generate documents for a verified assignment.
        Returns (docx_path, pdf_path) — empty strings on failure.
        """
        record = db.get(assignment_id)
        if not record:
            logger.error(f"DocGeneratorAgent: Assignment #{assignment_id} not found")
            return "", ""

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        title = record.get("title", "Assignment")
        subject = record.get("subject", "General")
        deadline = record.get("deadline", "N/A")
        answers: dict = record.get("answers", {})
        student_name = cfg.user_name

        safe_title = "".join(c for c in title if c.isalnum() or c in " _-")[:40].strip()
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        base_name = f"Assignment_{assignment_id}_{ts}"

        docx_path = str(OUTPUT_DIR / f"{base_name}.docx")
        pdf_path = str(OUTPUT_DIR / f"{base_name}.pdf")

        docx_ok = self._generate_docx(docx_path, title, subject, deadline, student_name, answers)
        pdf_ok = self._generate_pdf(pdf_path, title, subject, deadline, student_name, answers)

        if docx_ok and pdf_ok:
            db.update(assignment_id, docx_path=docx_path, pdf_path=pdf_path, status="pdf_ready")
            db.log(assignment_id, "doc_generated", f"DOCX+PDF: {base_name}")
            logger.info(f"DocGeneratorAgent done — {base_name}")
        else:
            db.set_status(assignment_id, "error", "Document generation failed")

        return (docx_path if docx_ok else ""), (pdf_path if pdf_ok else "")

    def _generate_docx(self, path: str, title: str, subject: str,
                       deadline: str, student: str, answers: dict) -> bool:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Meta info
            info = doc.add_paragraph()
            info.add_run(f"Subject: {subject}  |  Deadline: {deadline}  |  Student: {student}")
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Q&A
            for i, (question, answer) in enumerate(answers.items(), 1):
                q_para = doc.add_heading(f"Q{i}. {question}", level=2)
                doc.add_paragraph(answer)
                doc.add_paragraph()

            doc.save(path)
            logger.info(f"DOCX saved: {path}")
            return True
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return False

    def _generate_pdf(self, path: str, title: str, subject: str,
                      deadline: str, student: str, answers: dict) -> bool:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib import colors

            doc = SimpleDocTemplate(path, pagesize=A4,
                                    leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=16, spaceAfter=6)
            meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9,
                                        textColor=colors.grey, spaceAfter=12)
            q_style = ParagraphStyle("Q", parent=styles["Heading2"], fontSize=12, spaceAfter=4)
            a_style = ParagraphStyle("A", parent=styles["Normal"], fontSize=10,
                                     leading=14, spaceAfter=12)

            story = [
                Paragraph(title, title_style),
                Paragraph(f"Subject: {subject} | Deadline: {deadline} | Student: {student}", meta_style),
                Spacer(1, 0.3*cm),
            ]

            for i, (question, answer) in enumerate(answers.items(), 1):
                # Escape HTML special chars for reportlab
                safe_q = question.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_a = answer.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(f"Q{i}. {safe_q}", q_style))
                story.append(Paragraph(safe_a, a_style))
                story.append(Spacer(1, 0.2*cm))

            doc.build(story)
            logger.info(f"PDF saved: {path}")
            return True
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return False


doc_generator_agent = DocGeneratorAgent()
